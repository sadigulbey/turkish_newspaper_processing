import os, io
from google.cloud import vision
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches

def detect_text(path, count):
    """Detects text in the file."""
    from google.cloud import vision
    import io
    client = vision.ImageAnnotatorClient()

    with io.open(path, 'rb') as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations
    file = open("results/result-" + str(count) + ".txt", "a+")
    document = Document()
    
    p = document.add_paragraph("")
    for text in texts:
        print('\n"{}"'.format(text.description))
        if text.description.find("-") != -1:
            temp_content = text.description.replace("-\n","")
            temp_content = temp_content.replace("-", "")
            run = p.add_run(temp_content)
        else:
            run = p.add_run(temp_content +" ")
        #vertices = (['({},{})'.format(vertex.x, vertex.y)
                    #for vertex in text.bounding_poly.vertices])

        #print('bounds: {}'.format(','.join(vertices)))
    document.add_page_break()
    document.save("results/resultdoc.docx")
    file.close()


    if response.error.message:
        raise Exception(
            '{}\nFor more info on error messages, check: '
            'https://cloud.google.com/apis/design/errors'.format(
                response.error.message))





os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = r"Token.json"

client = vision.ImageAnnotatorClient()
pdfs = r"file.pdf"
pages = convert_from_path(pdfs)
img = pdfs.replace(".pdf","")

count = 1
for page in pages:
    jpeg_file = "images/" + img + "-" + str(count) + ".jpeg"
    page.save(jpeg_file, "JPEG")
    count += 1

for i in range(1, count):
    detect_text("images/" + img + "-" + str(i) + ".jpeg", i)

print("Completed")